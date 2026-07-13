"""
Generazione Word — sostituisce "Word Online" + export PDF di Power Automate.
Usa python-docx (nessuna dipendenza da Node/LibreOffice), quindi lo script
gira anche fuori da questo ambiente, sulla macchina dell'utente.
"""
import re
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH


def _add_markdown_paragraph(doc: Document, line: str):
    line = line.strip()
    if not line:
        return
    # bullet list
    if line.startswith("- ") or line.startswith("* "):
        doc.add_paragraph(line[2:].strip(), style="List Bullet")
        return
    # bold **text** -> single run bold (semplice, sufficiente per report generati da LLM)
    p = doc.add_paragraph()
    parts = re.split(r"(\*\*.*?\*\*)", line)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            run.bold = True
        else:
            run = p.add_run(part)


def build_docx(report_md: str, pratica: dict, output_path: str):
    doc = Document()

    # Frontespizio
    title = doc.add_heading("Relazione FIN-01 — Agente Finanziario EIAOS", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dati = pratica.get("dati_iniziali", {})
    meta.add_run(
        f"Pratica: {pratica['id']}\n"
        f"Strumento: {dati.get('strumento', 'n/d')}\n"
        f"ISIN/Ticker: {dati.get('isin_o_ticker', 'n/d')}\n"
        f"Data apertura: {pratica.get('data_apertura', 'n/d')}"
    ).italic = True

    disclaimer = doc.add_paragraph()
    disclaimer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    d_run = disclaimer.add_run(
        "Questa relazione è un supporto analitico alla decisione, basato su fonti pubbliche: "
        "non costituisce consulenza finanziaria personalizzata né garanzia di risultato. Le "
        "decisioni di impiego di capitale restano interamente a carico di chi le assume."
    )
    d_run.italic = True
    d_run.font.size = Pt(9)

    doc.add_page_break()

    # Corpo: parsing minimale del markdown restituito dal modello
    lines = report_md.splitlines()
    for line in lines:
        stripped = line.strip()
        if stripped.startswith("## "):
            doc.add_heading(stripped[3:].strip(), level=1)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:].strip(), level=2)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:].strip(), level=1)
        else:
            _add_markdown_paragraph(doc, line)

    doc.save(output_path)
    return output_path
